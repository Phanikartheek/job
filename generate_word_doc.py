"""
Generate Word Document from PROJECT_DOCUMENTATION.md
Run: python generate_word_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Helper: set paragraph shading ─────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Cover Page ─────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI-Powered Job Fraud Detection System')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Full Project Documentation')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Technology Stack: React + Flask + scikit-learn\n').bold = True
info.add_run('Project Type: Final Year Computer Science / AI Project\n')
info.add_run('Date: March 2026')

doc.add_page_break()

# ── Read markdown and convert ──────────────────────────────────
with open('PROJECT_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_lines = []

for line in lines:
    stripped = line.rstrip('\n')

    # Code block
    if stripped.startswith('```'):
        if in_code:
            # End code block
            code_text = '\n'.join(code_lines)
            cp = doc.add_paragraph()
            run = cp.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            shade_paragraph(cp, 'F4F4F4')
            in_code = False
            code_lines = []
        else:
            in_code = True
        continue

    if in_code:
        code_lines.append(stripped)
        continue

    # H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        p = doc.add_heading(stripped[2:], level=1)
        continue

    # H2
    if stripped.startswith('## '):
        p = doc.add_heading(stripped[3:], level=2)
        continue

    # H3
    if stripped.startswith('### '):
        p = doc.add_heading(stripped[4:], level=3)
        continue

    # H4
    if stripped.startswith('#### '):
        p = doc.add_heading(stripped[5:], level=4)
        continue

    # Table row
    if stripped.startswith('|') and '|' in stripped[1:]:
        if re.match(r'^\|[-| ]+\|$', stripped):
            continue  # separator row
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        # We'll just add as a plain paragraph to keep it simple
        p = doc.add_paragraph()
        p.add_run('  '.join(cells))
        continue

    # Bullet
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = stripped[2:]
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        p = doc.add_paragraph(text, style='List Bullet')
        continue

    # Horizontal rule
    if stripped.startswith('---'):
        p = doc.add_paragraph()
        p.add_run('─' * 80)
        continue

    # Bold inline
    if stripped:
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', stripped)
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        doc.add_paragraph()

# Save
out = 'Job_Fraud_Detection_Documentation.docx'
doc.save(out)
print(f'\n✅ Word document saved: {out}')
print(f'📁 Location: d:\\project 2\\job-main\\{out}')
